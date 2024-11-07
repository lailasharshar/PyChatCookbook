from docx.shared import Inches


# Function to add a page break if not the first recipe
def add_page_break_if_needed(doc, is_first):
    if not is_first:
        doc.add_page_break()


def add_recipe_to_document(doc, recipe_data, image_file):
    doc.add_heading(recipe_data['title'], level=1)
    doc.add_picture(image_file, width=Inches(5))
    # Add sub-items
    for sub_item in recipe_data.get('sub-items', []):
        doc.add_heading(sub_item['title'], level=2)

        # Ingredients
        doc.add_heading('Ingredients', level=3)
        for ingredient in sub_item.get('ingredients', []):
            amount = ingredient.get('amount', '')
            ingredient_name = ingredient.get('ingredient', '')
            description = ingredient.get('description', '')
            ingredient_line = f"- {amount} {ingredient_name} {description}".strip()
            doc.add_paragraph(ingredient_line)

        # Instructions
        doc.add_heading('Instructions', level=3)
        doc.add_paragraph(sub_item.get('instructions', ''))

    # Combination instructions
    combination_instructions = recipe_data.get('combination_instructions', [])
    if combination_instructions:
        doc.add_heading('Combination Instructions', level=2)
        for instruction in combination_instructions:
            doc.add_paragraph(instruction)
